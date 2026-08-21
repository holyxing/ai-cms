"""AI 文档导入任务 (P3.9.4+)

holy 反馈 2026-06-12 (#12096):
- 用户希望: 上传/粘贴 Word / PDF / 富文本, AI 自动解析成 HTML
- 图片自动上传到媒体库, HTML 里 <img src> 指向 MinIO 公开 URL
- AI 排版: 段落首行缩进 2em, 标题层级, 列表, 引用, 代码块

依据: docs/09-AI集成方案.md §4

3 个任务 (import_docx / import_pdf / import_paste_html) 共用
格式:
  1. 解析源文件 (docx → zip/xml, pdf → pymupdf, html → 正则)
  2. 提取图片 → 上传 MinIO → 拿到公开 URL
  3. 拼装 HTML (含上传后的图片 URL)
  4. (可选) AI 重新排版 (format_html 一样的 prompt)
  5. 写回 output.result_html

输出协议:
  {
    "result_html": str,           # 最终 HTML (含 MinIO 图片 URL)
    "description": str,           # "导入并排版了 X 个段落, 上传 Y 张图片"
    "diff_html": str,             # (同 result_html)
    "images_uploaded": list[dict], # [{local_name, url, size}, ...]
    "word_count": int,             # 中文字数
    "image_count": int,            # 处理的图片数
  }
"""
from __future__ import annotations

import io
import re
import uuid
import zipfile
from pathlib import Path
from typing import Optional
from xml.etree import ElementTree as ET

from loguru import logger
from sqlalchemy.ext.asyncio import AsyncSession

from app.agents.router import register
from app.models.ai_run import AIRun, AIRunStep
from app.services.llm.base import LLMMessage
from app.services.llm.factory import get_provider_for_user
from app.services.minio_client import presign_get_url, make_object_key, put_bytes


# === 任务类型 (注册名) ===
TASK_TYPES = ("import_docx", "import_pdf", "import_paste_html")

# 接收的 mime
MIME_MAP = {
    "import_docx": [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
        "application/zip",  # 有些客户端 mime 错
    ],
    "import_pdf": ["application/pdf"],
    "import_paste_html": ["text/html", "text/plain"],
}

# 排版 system prompt 已迁到 ai_prompts.key=import.layout（见 prompts/registry.py）
LAYOUT_PROMPT_FALLBACK = (
    "你是排版助手。将已解析的文档 HTML 做最小排版修正，输出必须看起来仍是同一篇。\n"
    "硬性要求:\n"
    "1. 禁止重写全文；保留原文内容与媒体标签\n"
    "2. 纯文本可用 <p> 分段，中文段落可加 style=\"text-indent: 2em;\"\n"
    "3. 列表用 <ul>/<ol>，标题用 <h2>/<h3>\n"
    "4. 直接输出 HTML，不要 markdown"
)


# ============ 工具函数 ============

def _sanitize_filename(name: str) -> str:
    """清掉文件名里危险字符"""
    name = re.sub(r"[^\w\u4e00-\u9fff\-_.]", "_", name)
    return name[:80] or f"image_{uuid.uuid4().hex[:8]}"


def _upload_image_to_minio(site_id: str, image_bytes: bytes, original_name: str, mime: str = "image/png") -> str:
    """上传图片到 MinIO + 返公开 URL

    上传后 LLM 拿到的 HTML 里的 <img> 替换为这个 URL, 用户编辑器里也是这个 URL
    """
    ext = Path(original_name).suffix.lower() or ".png"
    # mime 推断
    if mime == "image/jpeg" or ext == ".jpg":
        ext = ".jpg"
        mime = "image/jpeg"
    elif ext in (".png", ".gif", ".webp", ".svg", ".bmp", ".jpeg"):
        mime = f"image/{ext.lstrip('.')}" if ext != ".jpg" else "image/jpeg"
        if ext == ".jpeg":
            ext = ".jpg"
            mime = "image/jpeg"
    else:
        ext = ".png"
        mime = "image/png"
    object_key = make_object_key(site_id, ext)
    put_bytes(object_key, image_bytes, content_type=mime)
    url = presign_get_url(object_key, expires_seconds=3600 * 24 * 7)  # 7 天
    return url


# ============ 1. docx 解析 ============

def _parse_docx(file_bytes: bytes, site_id: str) -> tuple[str, list[dict]]:
    """解析 docx → (html, uploaded_images)

    docx = zip, 里面 word/document.xml 是文本, word/media/ 是图片
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(file_bytes))
    image_map: dict[str, str] = {}  # rid → minio_url
    uploaded: list[dict] = []

    # 1) 收集所有 inline 图片 (按 rId 索引)
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            try:
                image_bytes = rel.target_part.blob
                # mime 推断
                ext = Path(rel.target_ref).suffix.lower() or ".png"
                mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                            ".gif": "image/gif", ".webp": "image/webp", ".svg": "image/svg+xml"}
                mime = mime_map.get(ext, "image/png")
                url = _upload_image_to_minio(site_id, image_bytes, Path(rel.target_ref).name, mime)
                image_map[rel_id] = url
                uploaded.append({"local_name": Path(rel.target_ref).name, "url": url, "size": len(image_bytes)})
            except Exception as e:
                logger.warning(f"docx 图片上传失败: {e}")

    # 2) 遍历段落, 拼 HTML
    html_parts: list[str] = []
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower() if para.style else ""
        para_html = _render_paragraph_html(para, image_map)

        if "title" in style_name or "标题" in style_name or "heading 1" in style_name:
            html_parts.append(f"<h1>{para_html}</h1>")
        elif "heading 2" in style_name or "标题 2" in style_name:
            html_parts.append(f"<h2>{para_html}</h2>")
        elif "heading 3" in style_name or "标题 3" in style_name:
            html_parts.append(f"<h3>{para_html}</h3>")
        elif "heading 4" in style_name or "标题 4" in style_name:
            html_parts.append(f"<h4>{para_html}</h4>")
        elif "list" in style_name or "列表" in style_name:
            html_parts.append(f"<ul><li>{para_html}</li></ul>")
        elif para.text.strip() or para_html.strip():
            html_parts.append(f"<p>{para_html}</p>")
        # 空段跳过

    # 3) 表格
    # P3.9.6+ (holy 反馈 #12661): row.cells 返回 _Cell 对象, 不是 Paragraph — 不能直接 .runs
    # 正确做法: 对每个 cell, 遍历 cell.paragraphs (每个段落调用 _render_paragraph_html)
    for table in doc.tables:
        rows_html = []
        for row in table.rows:
            cells_html: list[str] = []
            for cell in row.cells:
                # 每个 cell 可能有多个段落 (一行中含换行)
                inner_parts = [
                    _render_paragraph_html(p, image_map) for p in cell.paragraphs if p.text.strip()
                ]
                cell_content = "".join(inner_parts) if inner_parts else ""
                cells_html.append(f"<td>{cell_content}</td>")
            rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
        if rows_html:
            html_parts.append(f"<table border='1'>{''.join(rows_html)}</table>")

    html = "\n".join(html_parts)
    return html, uploaded


def _render_paragraph_html(para, image_map: dict[str, str]) -> str:
    """docx 段落 → HTML 片段 (含 <img>)"""
    out: list[str] = []
    for run in para.runs:
        # 文本
        txt = run.text
        if not txt:
            # 检查 run 里有没有 inline 图片 (blip embed)
            blips = run._element.findall(".//" + "{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
            for blip in blips:
                embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if embed and embed in image_map:
                    out.append(f'<img src="{image_map[embed]}" alt="" />')
            continue
        # 简单样式
        if run.bold and run.italic:
            txt = f"<strong><em>{txt}</em></strong>"
        elif run.bold:
            txt = f"<strong>{txt}</strong>"
        elif run.italic:
            txt = f"<em>{txt}</em>"
        out.append(txt)
        # 行内图片
        blips = run._element.findall(".//" + "{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
        for blip in blips:
            embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if embed and embed in image_map:
                out.append(f'<img src="{image_map[embed]}" alt="" />')
    return "".join(out)


# ============ 2. PDF 解析 ============

def _parse_pdf(file_bytes: bytes, site_id: str) -> tuple[str, list[dict]]:
    """PDF → (html, uploaded_images)

    每页: 文本 + 图片 (按位置穿插)
    """
    import fitz  # pymupdf

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    uploaded: list[dict] = []
    html_parts: list[str] = []

    for page_idx, page in enumerate(doc):
        # 1) 文本块
        page_html_parts: list[str] = []
        text_blocks = page.get_text("dict")["blocks"]
        for block in text_blocks:
            if block.get("type") == 0:  # 文本块
                # 把 lines/spans 拼成一段
                lines_text = []
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if not text.strip():
                            continue
                        font_size = span.get("size", 12)
                        if font_size >= 24:
                            lines_text.append(f"<h1>{text}</h1>")
                        elif font_size >= 19:
                            lines_text.append(f"<h2>{text}</h2>")
                        elif font_size >= 16:
                            lines_text.append(f"<h3>{text}</h3>")
                        else:
                            lines_text.append(text)
                if lines_text:
                    page_html_parts.append("<p>" + "<br/>".join(lines_text) + "</p>")
        # 2) 图片 (提取每页所有图片)
        images = page.get_images(full=True)
        for img_idx, img in enumerate(images):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                ext = base_image["ext"]
                mime_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                            "gif": "image/gif", "webp": "image/webp"}
                mime = mime_map.get(ext, "image/png")
                fname = f"pdf_p{page_idx+1}_{img_idx+1}.{ext}"
                url = _upload_image_to_minio(site_id, image_bytes, fname, mime)
                uploaded.append({"local_name": fname, "url": url, "size": len(image_bytes)})
                # 插在文本末尾 (无法精确位置, 简化)
                page_html_parts.append(f'<p><img src="{url}" alt="" /></p>')
            except Exception as e:
                logger.warning(f"PDF 图片提取失败 p{page_idx+1} #{img_idx+1}: {e}")
        html_parts.extend(page_html_parts)

    doc.close()
    return "\n".join(html_parts), uploaded


# ============ 3. HTML/纯文本解析 (粘贴) ============

def _parse_html_or_text(content: str, site_id: str) -> tuple[str, list[dict]]:
    """粘贴的 HTML/纯文本 → (html, uploaded_images)

    - data:image/base64 内嵌图 → 上传 MinIO → 替换为公开 URL
    - 外部 img src (http) → 视情况: 可下到本地 (后续) / 或保留远程
    - 纯文本 → 简单换行 → <p>
    """
    uploaded: list[dict] = []

    # 1) base64 内嵌图 (Word/WPS 粘贴常见)
    def replace_data_uri(m: re.Match) -> str:
        mime = m.group(1) or "image/png"
        b64 = m.group(2)
        try:
            import base64
            image_bytes = base64.b64decode(b64)
            ext = {"image/png": ".png", "image/jpeg": ".jpg", "image/jpg": ".jpg",
                   "image/gif": ".gif", "image/webp": ".webp"}.get(mime, ".png")
            url = _upload_image_to_minio(site_id, image_bytes, f"pasted_{len(uploaded)+1}{ext}", mime)
            uploaded.append({"local_name": f"pasted_{len(uploaded)+1}{ext}", "url": url, "size": len(image_bytes)})
            return f'src="{url}"'
        except Exception as e:
            logger.warning(f"base64 图片解析失败: {e}")
            return m.group(0)

    content = re.sub(
        r'src=["\']data:([^;,]*?);base64,([^"\']+)["\']',
        replace_data_uri,
        content,
    )

    # 2) 纯文本 (无 HTML 标签) → 简单包 <p>
    if "<" not in content or not re.search(r"<\w+", content):
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]
        content = "\n".join(f"<p>{p}</p>" for p in paragraphs) if paragraphs else f"<p>{content}</p>"

    return content, uploaded


# ============ 注册 3 个任务 ============

@register("import_docx")
async def ai_import_docx(run, db, provider_model=None):
    return await _import_impl(run, db, provider_model, task_type="import_docx")


@register("import_pdf")
async def ai_import_pdf(run, db, provider_model=None):
    return await _import_impl(run, db, provider_model, task_type="import_pdf")


@register("import_paste_html")
async def ai_import_paste_html(run, db, provider_model=None):
    return await _import_impl(run, db, provider_model, task_type="import_paste_html")


# ============ 通用实现 ============

async def _import_impl(run: AIRun, db: AsyncSession, provider_model, task_type: str):
    input_data = run.input or {}
    site_id = str(run.site_id)
    file_url = input_data.get("file_url")  # 已上传到 MinIO 的源文件 URL (docx/pdf 走)
    content = input_data.get("content")  # 直接传的 HTML/文本 (paste_html 走)
    apply_layout = input_data.get("apply_layout", True)  # 是否调 AI 排版

    gen_step = AIRunStep(run_id=run.id, step_name="generate", step_order=2, status="running")
    db.add(gen_step)
    await db.flush()

    # 1. 解析
    # file_url 是 /media/... 相对路径, 走 nginx 会 403 (MinIO 对象私有 + 没签名).
    # 容器内直接 boto3 get_object, 绕过签名问题.
    async def _fetch_source_bytes(path_or_url: str) -> bytes:
        """/media/... 相对路径 (可能带 ?X-Amz-... 签名) → boto3 直拉 MinIO; 绝对 URL → httpx 拉 (兜底)"""
        if path_or_url.startswith("/"):
            from app.core.config import get_settings
            cfg = get_settings()
            import boto3
            s3 = boto3.client(
                "s3",
                endpoint_url=f"http://{cfg.MINIO_ENDPOINT}",
                aws_access_key_id=cfg.MINIO_ACCESS_KEY,
                aws_secret_access_key=cfg.MINIO_SECRET_KEY,
                region_name="us-east-1",
            )
            # 去掉路径里的 query string (presigned ?X-Amz-...) 和 /media/ 前缀
            key = path_or_url[len("/media/"):]
            if "?" in key:
                key = key.split("?", 1)[0]
            obj = s3.get_object(Bucket=cfg.MINIO_BUCKET, Key=key)
            return obj["Body"].read()
        # 兜底: 绝对 URL 用 httpx (例如外部图)
        import httpx
        async with httpx.AsyncClient() as client:
            r = await client.get(path_or_url, timeout=60)
            r.raise_for_status()
            return r.content

    if task_type == "import_docx":
        file_bytes = await _fetch_source_bytes(file_url)
        html, uploaded = _parse_docx(file_bytes, site_id)
    elif task_type == "import_pdf":
        file_bytes = await _fetch_source_bytes(file_url)
        html, uploaded = _parse_pdf(file_bytes, site_id)
    elif task_type == "import_paste_html":
        html, uploaded = _parse_html_or_text(content or "", site_id)
    else:
        raise ValueError(f"未知 import 任务: {task_type}")

    # 2. 基础清理
    html = _basic_sanitize(html)

    # 3. (可选) AI 排版 - 走 stream() 接口 (跟 format_html 一致, 边推边记 token)
    if apply_layout and html:
        try:
            from app.agents.prompts import resolve_prompt_content
            layout_prompt = await resolve_prompt_content(
                db, "import.layout", fallback=LAYOUT_PROMPT_FALLBACK,
            )
            provider = get_provider_for_user(run.user_id, provider_model)
            if provider:
                messages = [
                    LLMMessage(role="system", content=layout_prompt),
                    LLMMessage(role="user", content=html[:30000]),
                ]
                model = run.model or "qwen2.5:1.5b"
                # 输出接近输入；默认 8k 仍可能截断长文
                layout_max_tokens = min(32_000, max(8_192, int(min(len(html), 30000) * 1.25) + 2_048))
                accumulated: list[str] = []
                total_prompt_tokens = 0
                total_completion_tokens = 0
                async for chunk in provider.stream(
                    messages, model=model, temperature=0.2, max_tokens=layout_max_tokens,
                ):
                    if chunk.delta:
                        accumulated.append(chunk.delta)
                        gen_step.delta = "".join(accumulated[-200:])
                    if chunk.prompt_tokens:
                        total_prompt_tokens = chunk.prompt_tokens
                    if chunk.completion_tokens:
                        total_completion_tokens = chunk.completion_tokens
                laid_out = "".join(accumulated).strip()
                # 提取 HTML (兼容 code fence)
                laid_out = re.sub(r"^```html\s*", "", laid_out)
                laid_out = re.sub(r"^```\s*", "", laid_out)
                laid_out = re.sub(r"```$", "", laid_out).strip()
                if "<" in laid_out and len(laid_out) > 50:
                    html = laid_out
                # 写 token 到 run (跟 _text_transform 一样)
                if total_prompt_tokens or total_completion_tokens:
                    run.prompt_tokens = (run.prompt_tokens or 0) + total_prompt_tokens
                    run.completion_tokens = (run.completion_tokens or 0) + total_completion_tokens
                    # 算 cost (各 provider 自带 get_cost, 有就调)
                    try:
                        from app.services.llm.factory import get_provider_for_user as _gp
                        prov_cfg = _gp(run.user_id, provider_model)
                        if prov_cfg and hasattr(prov_cfg, "get_cost"):
                            cost = prov_cfg.get_cost(
                                total_prompt_tokens, total_completion_tokens, model,
                            )
                            from decimal import Decimal
                            run.cost_usd = (Decimal(str(run.cost_usd or 0)) + cost)
                    except Exception:
                        pass
        except Exception as e:
            logger.warning(f"AI 排版失败 (跳过): {e}")

    # 4. 字数统计
    word_count = len(re.sub(r"<[^>]+>", "", html))
    image_count = len(uploaded)

    # 5. 描述
    desc = f"导入并排版了 {word_count} 字, 上传 {image_count} 张图片"
    if image_count == 0:
        desc = f"导入并排版了 {word_count} 字"

    output = {
        "result_html": html,
        "description": desc,
        "diff_html": html,
        "images_uploaded": uploaded,
        "word_count": word_count,
        "image_count": image_count,
    }
    gen_step.output = output
    gen_step.status = "success"
    gen_step.finished_at = __import__("datetime").datetime.now(__import__("datetime").timezone.utc)
    await db.flush()
    return output


def _basic_sanitize(html: str) -> str:
    """基础 XSS 清理"""
    html = re.sub(r"<script\b[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r'\s+on\w+\s*=\s*["\'][^"\']*["\']', "", html, flags=re.IGNORECASE)
    html = re.sub(r"javascript\s*:", "", html, flags=re.IGNORECASE)
    return html
